"""Document extractor for Google Drive files and downloads.

Extracts plain text from downloaded or exported bytes according to mimeType
or file extension (PDF, DOCX, XLSX, Google Docs/Sheets/Slides, CSV, TXT, Markdown).
"""

from __future__ import annotations

import io
import re
from typing import Any


class DocumentExtractor:
    """Extract clean readable text from various document byte formats."""

    @staticmethod
    def extract_from_bytes(
        file_bytes: bytes,
        mime_type: str,
        file_name: str = "",
        max_chars: int = 12000,
    ) -> str:
        if not file_bytes:
            return ""

        text = ""
        lower_name = file_name.lower()

        try:
            # Skip binary archives, images, videos
            if (
                mime_type in ["application/zip", "application/x-zip-compressed", "application/x-tar", "application/gzip"]
                or any(lower_name.endswith(ext) for ext in [".zip", ".tar", ".gz", ".rar", ".7z", ".iso", ".bin", ".exe"])
            ):
                return f"[Archive file: {file_name} (Contents not indexed)]"

            if mime_type.startswith("image/") or any(lower_name.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]):
                return f"[Image file: {file_name}]"

            # Plain text / CSV / Markdown / Code / JSON / HTML
            if (
                mime_type.startswith("text/")
                or mime_type in [
                    "application/json",
                    "application/xml",
                    "application/javascript",
                    "application/x-yaml",
                    "text/markdown",
                    "text/csv",
                ]
                or any(lower_name.endswith(ext) for ext in [".txt", ".md", ".csv", ".json", ".py", ".js", ".html", ".log", ".yaml", ".yml", ".env"])
            ):
                text = DocumentExtractor._decode_text(file_bytes)

            # PDF
            elif mime_type == "application/pdf" or lower_name.endswith(".pdf"):
                text = DocumentExtractor._extract_pdf(file_bytes)

            # DOCX
            elif (
                mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or lower_name.endswith(".docx")
            ):
                text = DocumentExtractor._extract_docx(file_bytes)

            # XLSX / XLS
            elif (
                mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                or mime_type == "application/vnd.ms-excel"
                or lower_name.endswith(".xlsx")
                or lower_name.endswith(".xls")
            ):
                text = DocumentExtractor._extract_xlsx(file_bytes)

            # Google Drive exports (plain / csv)
            elif "plain" in mime_type or "csv" in mime_type:
                text = DocumentExtractor._decode_text(file_bytes)

            else:
                text = DocumentExtractor._decode_text(file_bytes)

        except Exception as e:
            text = f"[Error extracting text from {file_name}: {e}]"

        text = DocumentExtractor._clean_text(text)
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[...Truncated: {len(text) - max_chars} characters omitted...]"

        return text

    @staticmethod
    def _decode_text(data: bytes) -> str:
        for encoding in ["utf-8", "latin-1", "cp1252", "utf-16"]:
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        try:
            from pypdf import PdfReader
            stream = io.BytesIO(data)
            reader = PdfReader(stream)
            pages_text = []
            for i, page in enumerate(reader.pages):
                page_content = page.extract_text() or ""
                if page_content.strip():
                    pages_text.append(f"--- Page {i+1} ---\n{page_content.strip()}")
            return "\n\n".join(pages_text)
        except Exception as e:
            return f"[PDF parsing error: {e}]"

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        try:
            import docx
            stream = io.BytesIO(data)
            doc = docx.Document(stream)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as e:
            return f"[DOCX parsing error: {e}]"

    @staticmethod
    def _extract_xlsx(data: bytes) -> str:
        try:
            import openpyxl
            stream = io.BytesIO(data)
            wb = openpyxl.load_workbook(stream, data_only=True)
            sheets_text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                    if row_vals:
                        rows.append(" | ".join(row_vals))
                if rows:
                    sheets_text.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
            return "\n\n".join(sheets_text)
        except Exception:
            return DocumentExtractor._decode_text(data)

    @staticmethod
    def _clean_text(text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
